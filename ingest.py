"""
Document ingestion: load, chunk, embed, push to PGVector.
Supported types: PDF, Word (.docx/.doc), plain text (.txt/.md), images (.png/.jpg/.jpeg/.webp).
"""

from __future__ import annotations

import io
import os
import sys
import tempfile
from pathlib import Path

from dotenv import load_dotenv
import psycopg2
from langchain_postgres.vectorstores import PGVector
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from embeddings_factory import get_embeddings

BASE_DIR = Path(__file__).resolve().parent

# ---------------------------------------------------------------------------
# Per-type loaders
# ---------------------------------------------------------------------------

def _load_pdf_bytes(file_bytes: bytes, filename: str) -> list[Document]:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
        f.write(file_bytes)
        tmp_path = f.name
    try:
        try:
            from langchain_community.document_loaders import PyMuPDFLoader
            docs = PyMuPDFLoader(str(tmp_path)).load()
        except (ImportError, Exception):
            docs = PyPDFLoader(str(tmp_path)).load()
        for d in docs:
            d.metadata["source"] = filename
        return docs
    finally:
        os.unlink(tmp_path)


def _load_docx_bytes(file_bytes: bytes, filename: str) -> list[Document]:
    import docx
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
        f.write(file_bytes)
        tmp_path = f.name
    try:
        doc = docx.Document(tmp_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not text.strip():
            raise ValueError("No text extracted from docx")
        return [Document(page_content=text, metadata={"source": filename})]
    finally:
        os.unlink(tmp_path)


def _load_doc_bytes(file_bytes: bytes, filename: str) -> list[Document]:
    try:
        import docx2txt
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as f:
            f.write(file_bytes)
            tmp_path = f.name
        try:
            text = docx2txt.process(tmp_path)
            if not text or not text.strip():
                raise ValueError("No text extracted from .doc")
            return [Document(page_content=text, metadata={"source": filename})]
        finally:
            os.unlink(tmp_path)
    except Exception as e:
        raise ValueError(f"Could not load .doc file: {e}") from e


def _load_text_bytes(file_bytes: bytes, filename: str) -> list[Document]:
    text = file_bytes.decode("utf-8", errors="replace")
    if not text.strip():
        raise ValueError("Text file is empty")
    return [Document(page_content=text, metadata={"source": filename})]


def _load_image_bytes(file_bytes: bytes, filename: str) -> list[Document]:
    try:
        import pytesseract
        from PIL import Image
        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image)
        if not text.strip():
            raise ValueError("No text found in image via OCR")
        return [Document(page_content=text, metadata={"source": filename})]
    except ImportError as exc:
        raise ValueError(
            "pytesseract / Pillow not installed. Install them to use image OCR."
        ) from exc


_LOADERS: dict[str, object] = {
    ".pdf":  _load_pdf_bytes,
    ".docx": _load_docx_bytes,
    ".doc":  _load_doc_bytes,
    ".txt":  _load_text_bytes,
    ".md":   _load_text_bytes,
    ".png":  _load_image_bytes,
    ".jpg":  _load_image_bytes,
    ".jpeg": _load_image_bytes,
    ".webp": _load_image_bytes,
}

ALLOWED_EXTENSIONS: set[str] = set(_LOADERS.keys())


def load_document_bytes(file_bytes: bytes, filename: str) -> list[Document]:
    ext = Path(filename).suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    return loader(file_bytes, filename)  # type: ignore[operator]


# ---------------------------------------------------------------------------
# Single-document ingest (used by upload endpoint)
# ---------------------------------------------------------------------------

def ingest_document_bytes(
    file_bytes: bytes,
    filename: str,
    session_id: str,
    doc_id: str,
) -> int:
    """
    Chunk and embed a single document, tagging every chunk with session_id and doc_id.
    Returns the number of chunks added.
    Raises ValueError on bad input, Exception on infrastructure errors.
    """
    load_dotenv(BASE_DIR / ".env")
    db_url = os.environ.get("DATABASE_URL")
    if not db_url:
        raise ValueError("DATABASE_URL environment variable is not set")

    docs = load_document_bytes(file_bytes, filename)
    for d in docs:
        d.metadata["session_id"] = session_id
        d.metadata["doc_id"] = doc_id

    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=150)
    splits = splitter.split_documents(docs)

    if not splits:
        raise ValueError("No chunks produced — document may be empty.")

    embeddings = get_embeddings()
    vectorstore = PGVector(
        embeddings=embeddings,
        collection_name="rag_docs",
        connection=db_url,
        use_jsonb=True,
    )
    vectorstore.add_documents(splits)
    return len(splits)


# ---------------------------------------------------------------------------
# Full session re-index (used by CLI / migrate)
# ---------------------------------------------------------------------------

def rebuild_vector_index(session_id: str) -> dict[str, int]:
    """
    Re-embed all documents for a session. Clears existing session vectors first.
    Returns {"pdf_count": n, "chunk_count": m}.
    """
    load_dotenv(BASE_DIR / ".env")
    db_url = os.environ.get("DATABASE_URL")
    if not db_url:
        raise ValueError("DATABASE_URL environment variable is not set")

    docs: list[Document] = []
    file_count = 0

    with psycopg2.connect(db_url) as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT filename, content FROM document_files WHERE session_id = %s",
                (session_id,),
            )
            rows = cur.fetchall()

    for filename, content in rows:
        if not content:
            continue
        file_count += 1
        try:
            loaded = load_document_bytes(bytes(content), filename)
            for d in loaded:
                d.metadata["session_id"] = session_id
            docs.extend(loaded)
        except Exception as e:
            print(f"Warning: could not load {filename}: {e}", file=sys.stderr)

    if not file_count:
        raise ValueError("No documents found for this session. Upload a file first.")
    if not docs:
        raise ValueError("No text could be extracted from any document.")

    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=150)
    splits = splitter.split_documents(docs)
    embeddings = get_embeddings()

    vectorstore = PGVector(
        embeddings=embeddings,
        collection_name="rag_docs",
        connection=db_url,
        use_jsonb=True,
    )

    # Clear old session vectors before re-indexing
    try:
        with psycopg2.connect(db_url) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    DELETE FROM langchain_pg_embedding
                    WHERE collection_id IN (
                        SELECT uuid FROM langchain_pg_collection WHERE name = 'rag_docs'
                    )
                    AND cmetadata->>'session_id' = %s
                    """,
                    (session_id,),
                )
            conn.commit()
    except Exception as e:
        print(f"Note: could not clear old session embeddings: {e}")

    vectorstore.add_documents(splits)
    return {"pdf_count": file_count, "chunk_count": len(splits)}


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: python ingest.py <session_id>", file=sys.stderr)
        sys.exit(1)
    session_id = sys.argv[1]
    try:
        out = rebuild_vector_index(session_id)
    except ValueError as e:
        print(str(e), file=sys.stderr)
        sys.exit(1)
    print(
        f"Ingested {out['chunk_count']} chunks from {out['pdf_count']} file(s) "
        f"into PGVector (session={session_id})"
    )


if __name__ == "__main__":
    main()
